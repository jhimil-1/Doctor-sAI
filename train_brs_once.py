"""
Improved BRS document processing using LangChain for better chunking
"""
import os
import logging
from typing import List, Dict
from docx import Document
import sys
import os
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from openai import OpenAI
from pinecone import Pinecone, ServerlessSpec
from pinecone.exceptions import PineconeApiException
import time

# LangChain imports
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.document_loaders import Docx2txtLoader, PyPDFLoader
from langchain.schema import Document as LangChainDocument

from config import get_settings

logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

settings = get_settings()


class ImprovedBRSProcessor:
    """Enhanced BRS processor using LangChain for intelligent chunking"""
    
    def __init__(self):
        self.openai_client = OpenAI(api_key=settings.openai_api_key)
        # Initialize Pinecone
        self.pinecone = Pinecone(api_key=settings.pinecone_api_key)
        self.index_name = settings.pinecone_index_name
        
        # Initialize LangChain text splitter with document-aware settings
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=settings.chunk_size,
            chunk_overlap=settings.chunk_overlap,
            length_function=len,
            separators=[
                "\n\n\n",  # Multiple newlines (major sections)
                "\n\n",    # Double newlines (paragraphs)
                "\n",      # Single newlines
                ". ",      # Sentences
                ", ",      # Clauses
                " ",       # Words
                ""         # Characters (fallback)
            ],
            is_separator_regex=False,
        )
    
    def load_document(self, file_path: str) -> List[LangChainDocument]:
        """
        Load DOCX using LangChain loader with better structure preservation
        """
        try:
            # Choose loader based on file extension
            lower_path = file_path.lower()
            if lower_path.endswith('.pdf'):
                loader = PyPDFLoader(file_path)
            else:
                loader = Docx2txtLoader(file_path)
            documents = loader.load()
            
            # Add source metadata
            for doc in documents:
                doc.metadata["source"] = "BRS V1.2"
                doc.metadata["file_path"] = file_path
            
            logger.info(f"Loaded document with {len(documents)} sections")
            return documents
            
        except Exception as e:
            logger.error(f"Failed to load document: {e}")
            raise
    
    def extract_structured_content(self, file_path: str) -> List[LangChainDocument]:
        """
        Extract content with document structure (headings, sections) preserved
        """
        try:
            # Fallback for PDFs: structured extraction is DOCX-only
            if file_path.lower().endswith('.pdf'):
                logger.warning("Structured extraction not supported for PDF; using simple loader instead")
                return self.load_document(file_path)
            doc = Document(file_path)
            documents = []
            current_section = {"title": "", "content": [], "level": 0}
            
            for para in doc.paragraphs:
                text = para.text.strip()
                if not text:
                    continue
                
                # Detect headings based on style
                style = para.style.name.lower()
                
                if 'heading' in style:
                    # Save previous section if it has content
                    if current_section["content"]:
                        section_text = "\n".join(current_section["content"])
                        documents.append(LangChainDocument(
                            page_content=section_text,
                            metadata={
                                "source": "BRS V1.2",
                                "section": current_section["title"],
                                "heading_level": current_section["level"],
                                "type": "section"
                            }
                        ))
                    
                    # Start new section
                    level = int(style.replace('heading', '').strip() or '1')
                    current_section = {
                        "title": text,
                        "content": [f"## {text}"],  # Add heading to content
                        "level": level
                    }
                else:
                    current_section["content"].append(text)
            
            # Add last section
            if current_section["content"]:
                section_text = "\n".join(current_section["content"])
                documents.append(LangChainDocument(
                    page_content=section_text,
                    metadata={
                        "source": "BRS V1.2",
                        "section": current_section["title"],
                        "heading_level": current_section["level"],
                        "type": "section"
                    }
                ))
            
            # Extract tables separately
            for idx, table in enumerate(doc.tables):
                table_text = self._extract_table_text(table)
                if table_text:
                    documents.append(LangChainDocument(
                        page_content=table_text,
                        metadata={
                            "source": "BRS V1.2",
                            "type": "table",
                            "table_index": idx
                        }
                    ))
            
            logger.info(f"Extracted {len(documents)} structured sections")
            return documents
            
        except Exception as e:
            logger.error(f"Failed to extract structured content: {e}")
            raise
    
    def _extract_table_text(self, table) -> str:
        """Extract and format table content"""
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):  # Only add non-empty rows
                rows.append(" | ".join(cells))
        return "\n".join(rows) if rows else ""
    
    def chunk_documents(self, documents: List[LangChainDocument]) -> List[Dict]:
        """
        Chunk documents using LangChain's intelligent splitter
        """
        # Split documents into chunks
        chunks = self.text_splitter.split_documents(documents)
        
        # Convert to our format with enhanced metadata
        processed_chunks = []
        for idx, chunk in enumerate(chunks):
            processed_chunks.append({
                "id": f"brs_chunk_{idx}",
                "text": chunk.page_content,
                "metadata": {
                    **chunk.metadata,
                    "chunk_id": idx,
                    "char_count": len(chunk.page_content),
                    "chunk_method": "langchain_recursive"
                }
            })
        
        logger.info(f"Created {len(processed_chunks)} chunks using LangChain")
        
        # Log chunk size statistics
        chunk_sizes = [len(c["text"]) for c in processed_chunks]
        logger.info(f"Chunk size stats - Min: {min(chunk_sizes)}, "
                   f"Max: {max(chunk_sizes)}, "
                   f"Avg: {sum(chunk_sizes)//len(chunk_sizes)}")
        
        return processed_chunks
    
    def get_embeddings(self, texts: List[str], batch_size: int = 100) -> List[List[float]]:
        """
        Generate embeddings with batching for rate limit handling
        """
        all_embeddings = []
        
        for i in range(0, len(texts), batch_size):
            batch = texts[i:i + batch_size]
            try:
                response = self.openai_client.embeddings.create(
                    model=settings.embedding_model,
                    input=batch
                )
                embeddings = [item.embedding for item in response.data]
                all_embeddings.extend(embeddings)
                logger.info(f"Generated embeddings for batch {i//batch_size + 1} "
                           f"({len(embeddings)} texts)")
                
                # Small delay to avoid rate limits
                if i + batch_size < len(texts):
                    time.sleep(0.5)
                    
            except Exception as e:
                logger.error(f"Failed to generate embeddings for batch {i//batch_size + 1}: {e}")
                raise
        
        return all_embeddings
    
    def create_pinecone_index(self):
        """Ensure Pinecone index exists (idempotent)"""
        try:
            # If index already exists, just reuse it
            try:
                self.pinecone.describe_index(self.index_name)
                logger.info(f"Index '{self.index_name}' already exists. Using existing index.")
                return
            except Exception:
                pass

            logger.info(f"Creating index: {self.index_name}")
            self.pinecone.create_index(
                name=self.index_name,
                dimension=3072,  # text-embedding-3-large
                metric='cosine',
                spec=ServerlessSpec(
                    cloud='aws',
                    region=settings.pinecone_environment
                )
            )
            logger.info("Index created successfully")
            time.sleep(5)  # Wait for index to be ready
            
        except PineconeApiException as e:
            # Handle race or pre-existing index
            if getattr(e, 'status', None) == 409 or 'ALREADY_EXISTS' in str(e):
                logger.info(f"Index '{self.index_name}' already exists. Proceeding.")
                return
            logger.error(f"Failed to create index: {e}")
            raise
    
    def upload_to_pinecone(self, chunks: List[Dict]):
        """Upload chunks with embeddings to Pinecone"""
        try:
            index = self.pinecone.Index(self.index_name)
            
            # Get all texts for embedding
            texts = [chunk["text"] for chunk in chunks]
            embeddings = self.get_embeddings(texts)
            
            # Upload in batches
            batch_size = 100
            for i in range(0, len(chunks), batch_size):
                batch_chunks = chunks[i:i + batch_size]
                batch_embeddings = embeddings[i:i + batch_size]
                
                vectors = []
                for j, chunk in enumerate(batch_chunks):
                    vectors.append({
                        "id": chunk["id"],
                        "values": batch_embeddings[j],
                        "metadata": {
                            **chunk["metadata"],
                            "text": chunk["text"][:1000]  # Limit metadata text size
                        }
                    })
                
                index.upsert(vectors=vectors)
                logger.info(f"Uploaded batch {i // batch_size + 1} ({len(vectors)} vectors)")
            
            # Verify upload
            stats = index.describe_index_stats()
            logger.info(f"Total vectors in index: {stats['total_vector_count']}")
            
        except Exception as e:
            logger.error(f"Failed to upload to Pinecone: {e}")
            raise
    
    def process_and_upload(self, docx_path: str, use_structured: bool = True):
        """
        Main pipeline with LangChain integration
        
        Args:
            docx_path: Path to BRS document
            use_structured: Whether to use structured extraction (recommended)
        """
        logger.info("=" * 80)
        logger.info("Starting Enhanced BRS Processing with LangChain")
        logger.info("=" * 80)
        
        # Step 1: Load document
        logger.info("Step 1: Loading document...")
        if use_structured:
            documents = self.extract_structured_content(docx_path)
        else:
            documents = self.load_document(docx_path)
        
        # Step 2: Chunk with LangChain
        logger.info("Step 2: Chunking with LangChain...")
        chunks = self.chunk_documents(documents)
        
        # Step 3: Create index
        logger.info("Step 3: Setting up Pinecone...")
        self.create_pinecone_index()
        
        # Step 4: Upload
        logger.info("Step 4: Uploading to Pinecone...")
        self.upload_to_pinecone(chunks)
        
        logger.info("=" * 80)
        logger.info("✅ Processing Complete!")
        logger.info(f"Total chunks: {len(chunks)}")
        logger.info("=" * 80)


def main():
    """Run the improved processing pipeline"""
    BRS_DOCX_PATH = "C:/Users/nhz/Desktop/GenCare Assistant/03cb46934164321f675385fb74ac1bed.pdf"
    
    if not os.path.exists(BRS_DOCX_PATH):
        logger.error(f"BRS document not found at: {BRS_DOCX_PATH}")
        return
    
    try:
        processor = ImprovedBRSProcessor()
        
        # Use structured extraction for better results
        processor.process_and_upload(BRS_DOCX_PATH, use_structured=True)
        
        logger.info("✅ Training completed successfully!")
        
    except Exception as e:
        logger.error(f"❌ Training failed: {e}")
        raise


if __name__ == "__main__":
    main()